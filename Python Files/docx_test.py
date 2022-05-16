import docx
from docx.enum.text import WD_COLOR_INDEX
from docx_run import *

docx = docx.Document('RiskTest.docx')

run = isolate_run(docx.paragraphs[0], 5, 10)

run.font.highlight_color = WD_COLOR_INDEX.YELLOW

# for paragraph in docx.paragraphs:
#     for run in paragraph.runs:
#         print(type(run))

docx.save('Output.docx')