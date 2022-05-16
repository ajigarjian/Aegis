#TO DO: Add page 4 - 7 from OFRO and continue testing each page. Once done, look up how to modify word doc to add in comments or highlight the words.

#Importing packages

import sys

from pkg_resources import to_filename
sys.path.append('../bin/') 

import spacy
import docx
import os
import itertools
import copy
from spacy.matcher import Matcher
from docx_run import *
from docx.enum.text import WD_COLOR_INDEX

#Create the nlp object
nlp = spacy.load("en_core_web_sm")

# Intake a word doc into a list of nlp docs
docx = docx.Document('Sanitized Document.docx')
count = 0
docs = {}

for paragraph in docx.paragraphs:
    docs[count] = (nlp(paragraph.text))
    count += 1

# print(docs)

# Initialize the matcher with the shared vocab
matcher = Matcher(nlp.vocab)

#region
# Creating patterns for the each module in the OFRO list
pattern1_1 = [{"LOWER" : {"IN" : ["achievable", "attainable", "feasible"]}}]
pattern1_2 = [{"LEMMA" : "appropriate", "POS" : "ADJ"}]
pattern1_3 = [{"LEMMA" : "correct", "POS" : "ADJ"}]
pattern1_4 = [{"TEXT" : "PwC"}, {"LEMMA" : {"IN" : ["estimate", "forecast", "projection", "view"]}}]
pattern1_5 = [{"LOWER" : "suitable"}]
pattern1_6 = [{"LOWER" : "valid"}]

pattern2_1 = [{"LEMMA": "accurate", "OP": "?"}, {"TEXT": "accurately", "OP": "?"}]
pattern2_2 = [{"LEMMA": "all"}]
pattern2_3 = [{"LEMMA": "any"}]
pattern2_4 = [{"LEMMA": "complete", "OP": "?", "POS" : "ADJ"}, {"TEXT": "completely", "OP": "?"}]
pattern2_5 = [{"LEMMA": "comprehensive", "OP": "?"}, {"TEXT": "comprehensively", "OP": "?"}]
pattern2_6 = [{"LEMMA": "every"}]
pattern2_7 = [{"LEMMA": "full", "OP": "?", "POS" : "ADJ"}, {"TEXT": "fully", "OP": "?", "POS" : "ADV"}]
pattern2_8_1 = [{"LEMMA": "identify"}, {"POS": "ADP", "OP": "?"}, {"TEXT": "all"}]
pattern2_8_2 = [{"TEXT": "identification"}, {"POS": "ADP", "OP": "?"}, {"TEXT": "all"}]
pattern2_9_1 = [{"TEXT": "in"}, {"LEMMA": "accordance"}, {"TEXT": "with"}]
pattern2_9_2 = [{"LEMMA": "accord"}, {"TEXT": "to"}]
pattern2_10 = [{"LEMMA": "include"}, {"TEXT": "but"}, {"TEXT": "is", "OP": "?"}, {"TEXT": "not"}, {"LEMMA": "limit"}, {"POS": "ADP"}]
pattern2_11 = [{"LEMMA": {"IN": ["look", "examine", "study", "inspect", "scan", "scrutinize", "consider", "observe", "analyze", "review"]}}, {"POS": "ADP", "OP": "?"}, 
{"LEMMA": {"IN": ["all", "any", "every"]}}, {"LEMMA": {"IN": ["aspect", "angle", "feature", "facet", "particular", "detail", "view"]}}]
pattern2_12 = [{"LEMMA": "robust"}]

pattern3_1 = [{"LEMMA": "advocate"}]
pattern3_2 = [{"LEMMA": {"IN": ["collaborate", "collaboration"]}}]
pattern3_3 = [{"LEMMA": {"IN": ["coordinate", "coordination"]}}]
pattern3_4 = [{"LEMMA": "conclude"}]
pattern3_5 = [{"LEMMA": "determine"}]
pattern3_6 = [{"LEMMA": "endorse"}]
pattern3_7 = [{"LEMMA": "promote"}]
pattern3_8 = [{"LEMMA": "select"}]

pattern4_1 = [{"LOWER" : "agreed"}, {"LOWER" : "upon"}, {"LOWER" : "procedures"}]
pattern4_2 = [{"LEMMA": {"IN": ["attest", "attestation"]}}]
pattern4_3 = [{"LEMMA": {"IN": ["assure", "assurance"]}}]
pattern4_4 = [{"LEMMA": {"IN": ["compile", "compilation"]}}]
pattern4_5 = [{"LEMMA": "audit"}]
pattern4_6 = [{"LEMMA": "certify"}]
pattern4_7 = [{"LEMMA": {"IN": ["examine", "examination"]}}]
pattern4_8 = [{"LEMMA": "review"}]

pattern5_1 = [{"LEMMA": {"IN": ["approve", "approval"]}}]
pattern5_2 = [{"LEMMA": {"IN": ["assure", "assurance"]}}]
pattern5_3 = [{"LEMMA": "concur"}]
pattern5_4 = [{"LEMMA": "ensure"}]
pattern5_5 = [{"LEMMA": "guarantee"}]
pattern5_6 = [{"LEMMA": "insure"}]
pattern5_7 = [{"LEMMA": "promise"}]

pattern6_1 = [{"LEMMA" : {"IN" : ["ascertain", "know"]}}, {"LOWER": "your"}, {"LOWER" : {"IN" : ["needs", "requirements", "wants", "obligations", "demands", "obligations"]}}]

#pattern7_1 about red/yellow/green traffic lights
#pattern7_2 about low/medium/high traffic lights

pattern8_1_0 = [{"LOWER" : "industry", "OP" : "?"}, {"LOWER" : "best"}]
pattern8_1_1 = [{"LOWER" : "best"}, {"LEMMA" : {"IN" : ["possible", "practice"]}}]
pattern8_1_2 = [{"LOWER" : "best"}, {"LOWER" : "in"}, {"LOWER" : "class"}]
pattern8_2 = [{"LOWER" : {"IN" : ["cutting", "bleeding"]}}, {"LOWER" : "edge"}]
pattern8_3 = [{"LEMMA" : "very", "OP" : "?"}, {"LOWER" : "highest"}, {"LEMMA" : {"IN" : ["standard", "guideline", "benchmark", "measure"]}}]
pattern8_4 = [{"LOWER" : "ideal", "POS" : "ADJ"}]
pattern8_4_1 = [{"LEMMA" : "incomparable"}]
pattern8_5 = [{"LOWER" : "unequalled"}]
pattern8_5_1 = [{"LOWER" : {"IN" : ["efficient", "efficiency", "efficiently"]}}]
pattern8_6 = [{"LEMMA" : {"IN" : ["maximize", "maximum"]}}]
pattern8_7 = [{"LEMMA" : {"IN" : ["minimize", "minimum"]}}]
pattern8_8 = [{"LEMMA" : "normal"}]
pattern8_9 = [{"LEMMA" : "optimal"}]
pattern8_10 = [{"LEMMA" : "optimize"}]
pattern8_11 = [{"LOWER" : "state"}, {"LOWER" : "of"}, {"LOWER" : "the"}, {"LOWER" : "art"}]
pattern8_12 = [{"LOWER" : "sustainable"}]
pattern8_13 = [{"LEMMA" : "most", "OP" : "?"}, {"LOWER" : "tax"}, {"LOWER" : {"IN" : ["efficient", "efficiency", "scheme"]}}]
pattern8_14 = [{"LOWER" : "tax"}, {"IS_PUNCT": True, "OP" : "?"}, {"LEMMA" : {"IN" : ["minimize", "minimization"]}}]
pattern8_15 = [{"LOWER" : "world"}, {"LOWER" : "class"}]

pattern9_1 = [{"LOWER" : "earnings"}, {"LOWER" : "per"}, {"LOWER" : "share"}]
pattern9_2 = [{"LEMMA" : "shareholder"}]
pattern9_3 = [{"LOWER" : "value", "POS" : "NOUN"}]

pattern10_1 = [{"LEMMA" : "establish"}]
pattern10_2 = [{"LEMMA" : "implement"}]

pattern11 = [{"LEMMA" : "expert"}]

pattern12 = [{"LEMMA" : "extensive"}]

pattern13_1 = [{"LEMMA" : "high"}, {"IS_PUNCT": True, "OP" : "?"}, {"LEMMA" : "level"}]
pattern13_2 = [{"LOWER" : "detailed"}]

pattern14_1 = [{"LOWER" : {"IN" : ["immaterial", "immateriality"]}}]
pattern14_2 = [{"LOWER" : {"IN" : ["material", "materiality"]}}]

pattern15_1 = [{"LEMMA" : {"IN" : ["immediate", "immediately"]}}]
pattern15_2 = [{"LOWER" : "time"}, {"LOWER" : "is"}, {"LOWER" : "of"}, {"LOWER" : "the"}, {"LOWER" : "essence"}]
pattern15_2_1 = [{"LOWER" : "as"}, {"LOWER" : "soon"}, {"LOWER" : "as"}, {"LOWER" : "possible"}]
pattern15_3 = [{"LEMMA" : {"IN" : ["urgent", "urgently"]}}]

pattern16 = [{"LEMMA" : {"IN" : ["indemnify", "indemnity", "indemnification"]}}]

pattern17_1 = [{"LEMMA" : "invest"}, {"LOWER" : "in", "OP" : "?"}]
pattern17_2 = [{"LEMMA" : "investment"}]

pattern18_1 = [{"LEMMA" : {"IN" : ["know", "understand", "recognize"]}}, {"LOWER" : "your"}, {"LOWER" : {"IN" : ["needs", "requirements", "obligations", "necessities", "wants", "demands"]}}]
pattern18_2 = [{"LOWER" : "to"}, {"LOWER" : {"IN" : ["meet", "fulfill", "satisfy", "fill", "match"]}}, {"LOWER" : "all"}, {"LOWER" : "your"}, {"LOWER" : {"IN" : ["needs", "requirements", "obligations", "necessities", "wants", "demands"]}}]
pattern18_3 = [{"LOWER" : "to"}, {"LOWER" : {"IN" : ["your", "client's"]}}, {"LOWER" : "satisfaction"}]
pattern18_4 = [{"LEMMA" : {"IN" : ["satisfy", "satisfaction", "satisfactory"]}}]

pattern19 = [{"LOWER" : "legal"}, {"LOWER" : "document"}, {"LEMMA": {"IN": ["examination", "study", "inspection", "scan", "scrutinization", "consideration", "observation", "analysis", "review"]}}]

pattern20 = [{"LOWER" : "must"}]

pattern21 = [{"LEMMA" : {"IN" : ["negotiate", "negotiation"]}}]

pattern22 = [{"LOWER" : "next"}, {"LEMMA" : {"IN" : ["gen", "generation"]}}]

pattern23 = [{"LEMMA" : "opinion"}]

pattern24 = [{"LEMMA" : {"IN" : ["optimum", "optimal"]}}, {"LEMMA" : {"IN" : ["solution", "result", "answer", "conclusion", "finding"]}}]

pattern25_1 = [{"LEMMA" : {"IN" : ["partner", "partnership"]}}, {"LOWER" : "with"}]
pattern25_2 = [{"LEMMA" : "work"}, {"LOWER" : "jointly"}, {"LOWER" : "with"}]

pattern26 = [{"LEMMA" : "product"}]

pattern27 = [{"LOWER" : {"IN" : ["reasonable", "reasonably"]}}]

pattern28_1 = [{"LOWER" : "should"}]
pattern28_2 = [{"LEMMA" : {"IN" : ["will", "is"]}}]
pattern28_3 = [{"LOWER" : "more"}, {"LOWER" : "likely"}, {"LOWER" : "than"}, {"LOWER" : "not"}]

pattern29_1 = [{"LOWER" : "software"}, {"LEMMA" : "technology"}]
pattern29_2 = [{"LOWER" : {"IN" : ["technology", "technological"]}}, {"LEMMA" : "tool"}]

pattern30 = [{"LEMMA" : "support"}]

pattern31 = [{"LEMMA" : "tax"}]

pattern32 = [{"LEMMA" : "turnkey"}]

pattern33_1 = [{"LEMMA" : "validate"}]
pattern33_2 = [{"LEMMA" : "verify"}]

#creating a dictionary that matches each pattern to its suggestion

suggestions = {}
suggestions["pattern1_1"] = """This word should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated. 
Our role is not to offer concrete conclusions/solutions/views, rather it is to do the analysis and allow the client to use the information presented to draw their own conclusions. 
PwC's role should be advisory in nature, we should avoid attest-type terms."""
suggestions["pattern1_2"] = """This word should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated. 
Our role is not to offer concrete conclusions/solutions/views, rather it is to do the analysis and allow the client to use the information presented to draw their own conclusions. 
PwC's role should be advisory in nature, we should avoid attest-type terms."""
suggestions["pattern1_3"] = """This word should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated. 
Our role is not to offer concrete conclusions/solutions/views, rather it is to do the analysis and allow the client to use the information presented to draw their own conclusions. 
PwC's role should be advisory in nature, we should avoid attest-type terms."""
suggestions["pattern1_4"] = """This word should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated. 
Our role is not to offer concrete conclusions/solutions/views, rather it is to do the analysis and allow the client to use the information presented to draw their own conclusions. 
PwC's role should be advisory in nature, we should avoid attest-type terms.
Unless you are providing Deals services, do not prepare entity level PFI and do not attribute entity level PFI to PwC (e.g., do not label entity level analysis as PwC Estimate, PwC Projections, PwC Bse Case, PwC view, etc., but clearly attribute the original PFI to client/target)."""
suggestions["pattern1_5"] = """This word should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated. 
Our role is not to offer concrete conclusions/solutions/views, rather it is to do the analysis and allow the client to use the information presented to draw their own conclusions. 
PwC's role should be advisory in nature, we should avoid attest-type terms."""
suggestions["pattern1_6"] = """This word should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated. 
Our role is not to offer concrete conclusions/solutions/views, rather it is to do the analysis and allow the client to use the information presented to draw their own conclusions. 
PwC's role should be advisory in nature, we should avoid attest-type terms."""

suggestions["pattern2_1"] = """Care must be taken to confirm we are not using words that imply our scope was more broad or detailed than actually performed. We make no implied actual guarantees.
Avoid absolutes (i.e., "all") in context of providing advice. Similarly, different readers may have different interpretations of what those words mean. These words should not be used in relation to our services.
Can mean marked by richness and fullness; the issue is fullness suggests that it has everything."""
suggestions["pattern2_2"] = """Care must be taken to confirm we are not using words that imply our scope was more broad or detailed than actually performed. We make no implied actual guarantees.
Avoid absolutes (i.e., "all") in context of providing advice. Similarly, different readers may have different interpretations of what those words mean. These words should not be used in relation to our services.
Can mean marked by richness and fullness; the issue is fullness suggests that it has everything."""
suggestions["pattern2_3"] = """Care must be taken to confirm we are not using words that imply our scope was more broad or detailed than actually performed. We make no implied actual guarantees.
Avoid absolutes (i.e., "all") in context of providing advice. Similarly, different readers may have different interpretations of what those words mean. These words should not be used in relation to our services.
Can mean marked by richness and fullness; the issue is fullness suggests that it has everything."""
suggestions["pattern2_4"] = """Care must be taken to confirm we are not using words that imply our scope was more broad or detailed than actually performed. We make no implied actual guarantees.
Avoid absolutes (i.e., "all") in context of providing advice. Similarly, different readers may have different interpretations of what those words mean. These words should not be used in relation to our services.
Can mean marked by richness and fullness; the issue is fullness suggests that it has everything."""
suggestions["pattern2_5"] = """Care must be taken to confirm we are not using words that imply our scope was more broad or detailed than actually performed. We make no implied actual guarantees.
Avoid absolutes (i.e., "all") in context of providing advice. Similarly, different readers may have different interpretations of what those words mean. These words should not be used in relation to our services.
Can mean marked by richness and fullness; the issue is fullness suggests that it has everything."""
suggestions["pattern2_6"] = """Care must be taken to confirm we are not using words that imply our scope was more broad or detailed than actually performed. We make no implied actual guarantees.
Avoid absolutes (i.e., "all") in context of providing advice. Similarly, different readers may have different interpretations of what those words mean. These words should not be used in relation to our services.
Can mean marked by richness and fullness; the issue is fullness suggests that it has everything."""
suggestions["pattern2_7"] = """Care must be taken to confirm we are not using words that imply our scope was more broad or detailed than actually performed. We make no implied actual guarantees.
Avoid absolutes (i.e., "all") in context of providing advice. Similarly, different readers may have different interpretations of what those words mean. These words should not be used in relation to our services.
Can mean marked by richness and fullness; the issue is fullness suggests that it has everything."""
suggestions["pattern2_8_1"] = """Care must be taken to confirm we are not using words that imply our scope was more broad or detailed than actually performed. We make no implied actual guarantees.
Avoid absolutes (i.e., "all") in context of providing advice. Similarly, different readers may have different interpretations of what those words mean. These words should not be used in relation to our services.
Can mean marked by richness and fullness; the issue is fullness suggests that it has everything."""
suggestions["pattern2_8_2"] = """Care must be taken to confirm we are not using words that imply our scope was more broad or detailed than actually performed. We make no implied actual guarantees.
Avoid absolutes (i.e., "all") in context of providing advice. Similarly, different readers may have different interpretations of what those words mean. These words should not be used in relation to our services.
Can mean marked by richness and fullness; the issue is fullness suggests that it has everything."""
suggestions["pattern2_9_1"] = """Care must be taken to confirm we are not using words that imply our scope was more broad or detailed than actually performed. We make no implied actual guarantees.
Avoid absolutes (i.e., "all") in context of providing advice. Similarly, different readers may have different interpretations of what those words mean. These words should not be used in relation to our services.
Can mean marked by richness and fullness; the issue is fullness suggests that it has everything."""
suggestions["pattern2_9_2"] = """Care must be taken to confirm we are not using words that imply our scope was more broad or detailed than actually performed. We make no implied actual guarantees.
Avoid absolutes (i.e., "all") in context of providing advice. Similarly, different readers may have different interpretations of what those words mean. These words should not be used in relation to our services.
Can mean marked by richness and fullness; the issue is fullness suggests that it has everything."""
suggestions["pattern2_10"] = """Care must be taken to confirm we are not using words that imply our scope was more broad or detailed than actually performed. We make no implied actual guarantees.
Avoid absolutes (i.e., "all") in context of providing advice. Similarly, different readers may have different interpretations of what those words mean. These words should not be used in relation to our services.
Can mean marked by richness and fullness; the issue is fullness suggests that it has everything."""
suggestions["pattern2_11"] = """Care must be taken to confirm we are not using words that imply our scope was more broad or detailed than actually performed. We make no implied actual guarantees.
Avoid absolutes (i.e., "all") in context of providing advice. Similarly, different readers may have different interpretations of what those words mean. These words should not be used in relation to our services.
Can mean marked by richness and fullness; the issue is fullness suggests that it has everything."""
suggestions["pattern2_12"] = """Care must be taken to confirm we are not using words that imply our scope was more broad or detailed than actually performed. We make no implied actual guarantees.
Avoid absolutes (i.e., "all") in context of providing advice. Similarly, different readers may have different interpretations of what those words mean. These words should not be used in relation to our services.
Can mean marked by richness and fullness; the issue is fullness suggests that it has everything."""

suggestions["pattern3_1"] = """We avoid the appearance of taking on client's management responsibilities."""
suggestions["pattern3_2"] = """We avoid the appearance of taking on client's management responsibilities. 'Collaborate/collaboration' should be used cautiously with Risk Management consultation and only in the development of unbranded deliverables. The use of 'collaborate/collaboration' may be permissible in the context of speaking about our existing/approved JBR entity relationships, including non-profit organizations."""
suggestions["pattern3_3"] = """We avoid the appearance of taking on client's management responsibilities. If we are using "coordinate", make sure the context of its use is clear that we are assisting management in its coordination and not taking on a management role."""
suggestions["pattern3_4"] = """We avoid the appearance of taking on client's management responsibilities."""
suggestions["pattern3_5"] = """We avoid the appearance of taking on client's management responsibilities."""
suggestions["pattern3_6"] = """We avoid the appearance of taking on client's management responsibilities."""
suggestions["pattern3_7"] = """We avoid the appearance of taking on client's management responsibilities."""
suggestions["pattern3_8"] = """We avoid the appearance of taking on client's management responsibilities."""

suggestions["pattern4_1"] = """Words defined in the professional standards should only be used in a manner consistent with that meaning or when performing those specific services.
The more appropriate word is often 'read'."""
suggestions["pattern4_2"] = """Words defined in the professional standards should only be used in a manner consistent with that meaning or when performing those specific services.
The more appropriate word is often 'read'.""" 
suggestions["pattern4_3"] = """Words defined in the professional standards should only be used in a manner consistent with that meaning or when performing those specific services.
The more appropriate word is often 'read'."""
suggestions["pattern4_4"] = """Words defined in the professional standards should only be used in a manner consistent with that meaning or when performing those specific services.
The more appropriate word is often 'read'."""
suggestions["pattern4_5"] = """Words defined in the professional standards should only be used in a manner consistent with that meaning or when performing those specific services.
The more appropriate word is often 'read'."""
suggestions["pattern4_6"] = """Words defined in the professional standards should only be used in a manner consistent with that meaning or when performing those specific services.
The more appropriate word is often 'read'."""
suggestions["pattern4_7"] = """Words defined in the professional standards should only be used in a manner consistent with that meaning or when performing those specific services.
The more appropriate word is often 'read'."""
suggestions["pattern4_8"] = """Words defined in the professional standards should only be used in a manner consistent with that meaning or when performing those specific services.
The more appropriate word is often 'read'. In the very limited circumstances where the word 'review' is used, it may be appropriate when 'reviewing' a client or target-provided document or schedule and the context is clear that the use of the world 'review' would not create an expectation that PwC's 'review' includes providing assurance over such document or schedule. Further, the extent of any review should be articulated to avoid a misunderstanding of the scope of our work."""

suggestions["pattern5_1"] = """This word should not be used in connection with our services. We make no implied/actual guarantees.
In addition, care must be taken to confirm we are not using words that give the impression of providing attest services, 'negative assurance' or implying our scope was more detailed than actually performed and/or providing more certainty than intended.
Use of alternative words should also include clarity or specificity as to what we will be assisting with. Further, if you are working with a restricted entity and following the 5 or 7 step advisory process be sure to follow guidance from the Independence Office for how to describe our involvement or services."""
suggestions["pattern5_2"] = """This word should not be used in connection with our services. We make no implied/actual guarantees.
In addition, care must be taken to confirm we are not using words that give the impression of providing attest services, 'negative assurance' or implying our scope was more detailed than actually performed and/or providing more certainty than intended.
Use of alternative words should also include clarity or specificity as to what we will be assisting with. Further, if you are working with a restricted entity and following the 5 or 7 step advisory process be sure to follow guidance from the Independence Office for how to describe our involvement or services."""
suggestions["pattern5_3"] = """This word should not be used in connection with our services. We make no implied/actual guarantees.
In addition, care must be taken to confirm we are not using words that give the impression of providing attest services, 'negative assurance' or implying our scope was more detailed than actually performed and/or providing more certainty than intended.
Use of alternative words should also include clarity or specificity as to what we will be assisting with. Further, if you are working with a restricted entity and following the 5 or 7 step advisory process be sure to follow guidance from the Independence Office for how to describe our involvement or services."""
suggestions["pattern5_4"] = """This word should not be used in connection with our services. We make no implied/actual guarantees.
In addition, care must be taken to confirm we are not using words that give the impression of providing attest services, 'negative assurance' or implying our scope was more detailed than actually performed and/or providing more certainty than intended.
Use of alternative words should also include clarity or specificity as to what we will be assisting with. Further, if you are working with a restricted entity and following the 5 or 7 step advisory process be sure to follow guidance from the Independence Office for how to describe our involvement or services."""
suggestions["pattern5_5"] = """This word should not be used in connection with our services. We make no implied/actual guarantees.
In addition, care must be taken to confirm we are not using words that give the impression of providing attest services, 'negative assurance' or implying our scope was more detailed than actually performed and/or providing more certainty than intended.
Use of alternative words should also include clarity or specificity as to what we will be assisting with. Further, if you are working with a restricted entity and following the 5 or 7 step advisory process be sure to follow guidance from the Independence Office for how to describe our involvement or services."""
suggestions["pattern5_6"] = """This word should not be used in connection with our services. We make no implied/actual guarantees.
In addition, care must be taken to confirm we are not using words that give the impression of providing attest services, 'negative assurance' or implying our scope was more detailed than actually performed and/or providing more certainty than intended.
Use of alternative words should also include clarity or specificity as to what we will be assisting with. Further, if you are working with a restricted entity and following the 5 or 7 step advisory process be sure to follow guidance from the Independence Office for how to describe our involvement or services."""
suggestions["pattern5_7"] = """This word should not be used in connection with our services. We make no implied/actual guarantees.
In addition, care must be taken to confirm we are not using words that give the impression of providing attest services, 'negative assurance' or implying our scope was more detailed than actually performed and/or providing more certainty than intended.
Use of alternative words should also include clarity or specificity as to what we will be assisting with. Further, if you are working with a restricted entity and following the 5 or 7 step advisory process be sure to follow guidance from the Independence Office for how to describe our involvement or services."""

suggestions["pattern6_1"] = """We advise/assist; we do not promise results."""

#pattern7_1 about red/yellow/green traffic lights
#pattern7_2 about low/medium/high traffic lights

suggestions["pattern8_1_0"] = """This word/phrase should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated. 
Avoid absolutes (i.e., "all"). Absolutes should be avoided in the context of providing advice. Similarly, different readers may have different interpretations of what those words mean.
These words should not be used in relation to our services; however, use of these words may be acceptable where we are simply recognizing that the client would like to 'maximize', 'minimize', 'optimize', for example, a process or system."""
suggestions["pattern8_1_1"] = """This word/phrase should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated. 
Avoid absolutes (i.e., "all"). Absolutes should be avoided in the context of providing advice. Similarly, different readers may have different interpretations of what those words mean.
These words should not be used in relation to our services; however, use of these words may be acceptable where we are simply recognizing that the client would like to 'maximize', 'minimize', 'optimize', for example, a process or system."""
suggestions["pattern8_1_2"] = """This word/phrase should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated. 
Avoid absolutes (i.e., "all"). Absolutes should be avoided in the context of providing advice. Similarly, different readers may have different interpretations of what those words mean.
These words should not be used in relation to our services; however, use of these words may be acceptable where we are simply recognizing that the client would like to 'maximize', 'minimize', 'optimize', for example, a process or system."""
suggestions["pattern8_2"] = """This phrase should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated."""
suggestions["pattern8_3"] = """This phrase should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated. 
Avoid absolutes (i.e., "all"). Absolutes should be avoided in the context of providing advice. Similarly, different readers may have different interpretations of what those words mean."""
suggestions["pattern8_4"] = """This word should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated. 
Avoid absolutes (i.e., "all"). Absolutes should be avoided in the context of providing advice. Similarly, different readers may have different interpretations of what those words mean."""
suggestions["pattern8_4_1"] = """This word should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated. 
Avoid absolutes (i.e., "all"). Absolutes should be avoided in the context of providing advice. Similarly, different readers may have different interpretations of what those words mean."""
suggestions["pattern8_5"] = """This word should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated. 
Avoid absolutes (i.e., "all"). Absolutes should be avoided in the context of providing advice. Similarly, different readers may have different interpretations of what those words mean."""
suggestions["pattern8_5_1"] = """This word should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated. 
Efficient is subjective and also may be viewed as not in keeping with our tax code of conduct. For example in some other countries the term "scheme" is used. In the US the phrase "tax scheme" has a negative connotation. The same may be said of "tax efficient". Clients should determine the tax consequences we calculate/describe/determine and if they are suitable. We should not state they are "efficient"."""
suggestions["pattern8_6"] = """This word should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated. 
Avoid absolutes (i.e., "all"). Absolutes should be avoided in the context of providing advice. Similarly, different readers may have different interpretations of what those words mean.
These words should not be used in relation to our services; however, use of these words may be acceptable where we are simply recognizing that the client would like to 'maximize', 'minimize', 'optimize', for example, a process or system."""
suggestions["pattern8_7"] = """This word should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated. 
Avoid absolutes (i.e., "all"). Absolutes should be avoided in the context of providing advice. Similarly, different readers may have different interpretations of what those words mean.
These words should not be used in relation to our services; however, use of these words may be acceptable where we are simply recognizing that the client would like to 'maximize', 'minimize', 'optimize', for example, a process or system."""
suggestions["pattern8_8"] = """This word should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated."""
suggestions["pattern8_9"] = """This word should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated. 
Avoid absolutes (i.e., "all"). Absolutes should be avoided in the context of providing advice. Similarly, different readers may have different interpretations of what those words mean.
These words should not be used in relation to our services; however, use of these words may be acceptable where we are simply recognizing that the client would like to 'maximize', 'minimize', 'optimize', for example, a process or system."""
suggestions["pattern8_10"] = """This word should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated. 
Avoid absolutes (i.e., "all"). Absolutes should be avoided in the context of providing advice. Similarly, different readers may have different interpretations of what those words mean.
These words should not be used in relation to our services; however, use of these words may be acceptable where we are simply recognizing that the client would like to 'maximize', 'minimize', 'optimize', for example, a process or system."""
suggestions["pattern8_11"] = """This phrase should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated."""
suggestions["pattern8_12"] = """This word should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated.
Sustainable may be permissible in the context of speaking about environmental aspirational goals, value chains (sustaining trust), and language relating to The New Equation (e.g., sustainable outcomes)."""
suggestions["pattern8_13"] = """This phrase should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated. 
Efficient is subjective and also may be viewed as not in keeping with our tax code of conduct. For example in some other countries the term "scheme" is used. In the US the phrase "tax scheme" has a negative connotation. The same may be said of "tax efficient". Clients should determine the tax consequences we calculate/describe/determine and if they are suitable. We should not state they are "efficient". 
Avoid absolutes (i.e., "all"). Absolutes should be avoided in the context of providing advice. Similarly, different readers may have different interpretations of what those words mean."""
suggestions["pattern8_14"] = """This phrase should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated. 
Efficient is subjective and also may be viewed as not in keeping with our tax code of conduct. For example in some other countries the term "scheme" is used. In the US the phrase "tax scheme" has a negative connotation. The same may be said of "tax efficient". Clients should determine the tax consequences we calculate/describe/determine and if they are suitable. We should not state they are "efficient". 
Avoid absolutes (i.e., "all"). Absolutes should be avoided in the context of providing advice. Similarly, different readers may have different interpretations of what those words mean."""
suggestions["pattern8_15"] = """This phrase should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated.
Avoid absolutes (i.e., "all"). Absolutes should be avoided in the context of providing advice. Similarly, different readers may have different interpretations of what those words mean."""

suggestions["pattern9_1"] = """Avoid this word/phrase. We do not give advice e.g., financial statements."""
suggestions["pattern9_2"] = """Avoid this word/phrase. We do not give advice e.g., financial statements."""
suggestions["pattern9_3"] = """Avoid this word/phrase. We do not give advice e.g., financial statements."""

suggestions["pattern10_1"] = """Establishing something implies we are performing management functions around the establishment (vs. the client being the decision maker.
Use of an alternative word such as "assist" should also include clarity or specificity as to what we will be assisting with."""
suggestions["pattern10_2"] = """Implementing something implies we are performing management functions around the implementation (vs. the client being the decision maker.
We advise a client and assist them with their implementation of the solution(s), etc., however they own the actual implementation. For example: We will assist client with its implementation of X.
Use of an alternative word such as "assist" should also include clarity or specificity as to what we will be assisting with."""

suggestions["pattern11"] = """Unsubstantiated; implies earned credential.
Expert may be used to identify an individual that has been appointed as an 'expert witness' in a court of law. Outside of this, "expert" should not be used."""

suggestions["pattern12"] = """The sufficiency of scope should be assessed by the client not by PwC."""

suggestions["pattern13_1"] = """Ambiguous - can mean different things to different people.
High-level - used the phrase cautiously and it should generally not be used casually without a more detailed description of our work. Engagement teams should confirm that there is sufficient further detail in the description of services to allow both the client and PwC to have a clear understanding of what work will be done and will not be done in the context of "high-level"."""
suggestions["pattern13_2"] = """Ambiguous - can mean different things to different people."""

suggestions["pattern14_1"] = """It is important to consider the context in which this word is being used. "Materiality" is an audit concept and generally should not be used in non audit related communications.
Using the word material or immaterial in the context of qualitative measures may be acceptable provided that the basis is documented within the communication. However, the word material/immaterial should not be used in the context of qualitative measures (e.g. account balances, proposed adjustments) unless defined by the client (e.g., if the client sets a materiality threshold for our reporting of tax exposures in a due diligence engagement).
Even the use of alternative words/phrases such as significant/insignificant/de minimis should be used cautiously, and used when the expectation/context is clear to all parties. Materiality should be assessed by the client and not the engagement team."""
suggestions["pattern14_2"] = """It is important to consider the context in which this word is being used. "Materiality" is an audit concept and generally should not be used in non audit related communications.
Using the word material or immaterial in the context of qualitative measures may be acceptable provided that the basis is documented within the communication. However, the word material/immaterial should not be used in the context of qualitative measures (e.g. account balances, proposed adjustments) unless defined by the client (e.g., if the client sets a materiality threshold for our reporting of tax exposures in a due diligence engagement).
Even the use of alternative words/phrases such as significant/insignificant/de minimis should be used cautiously, and used when the expectation/context is clear to all parties. Materiality should be assessed by the client and not the engagement team."""

suggestions["pattern15_1"] = """Immediately is an unreasonably high standard to achieve.
This word signifies that something should/can be done very quickly and perhaps implies lack of due care. The provision of all our services are dependent on the client and our team working together and agreeing to a delivery schedule. This word may be appropriate if not associated with our services."""
suggestions["pattern15_2"] = """This phrase signifies that something should/can be done very quickly and perhaps implies lack of due care. The provision of all our services are dependent on the client and our team working together and agreeing to a delivery schedule. This phrase may be appropriate if not associated with our services."""
suggestions["pattern15_2_1"] = """This phrase signifies that something should/can be done very quickly and perhaps implies lack of due care. The provision of all our services are dependent on the client and our team working together and agreeing to a delivery schedule. This phrase may be appropriate if not associated with our services."""
suggestions["pattern15_3"] = """This word signifies that something should/can be done very quickly and perhaps implies lack of due care. The provision of all our services are dependent on the client and our team working together and agreeing to a delivery schedule. This word may be appropriate if not associated with our services."""

suggestions["pattern16"] = """PwC generally does not indemnitfy its clients or third parties when engaged to perform professional services except as included in our standard contracting terms and conditions."""

suggestions["pattern17_1"] = """This word should not be used in connection with describing our relationship with the client and associated fee discounts. We cannot "invest in" clients from whom we must remain independent."""
suggestions["pattern17_2"] = """This word should not be used in connection with describing our relationship with the client and associated fee discounts. We cannot "invest in" clients from whom we must remain independent."""

suggestions["pattern18_1"] = """Ambiguous: Implies that we will carry on regardless of cost until the client is satisfied.
Our role is to advise/assist clients; we do not promise results. Client is responsible for assessing the adequacy of the scope and our work in relation to their needs."""
suggestions["pattern18_2"] = """Ambiguous: Implies that we will carry on regardless of cost until the client is satisfied.
Our role is to advise/assist clients; we do not promise results. Client is responsible for assessing the adequacy of the scope and our work in relation to their needs."""
suggestions["pattern18_3"] = """Ambiguous: Implies that we will carry on regardless of cost until the client is satisfied.
Our role is to advise/assist clients; we do not promise results. Client is responsible for assessing the adequacy of the scope and our work in relation to their needs."""
suggestions["pattern18_4"] = """Ambiguous: Implies that we will carry on regardless of cost until the client is satisfied.
Our role is to advise/assist clients; we do not promise results. Client is responsible for assessing the adequacy of the scope and our work in relation to their needs."""

suggestions["pattern19"] = """We are not permitted to engage in the practice of law. We can provide comments related to our area of subject matter specialization but should refrain from providing other comments or drafting suggested legal language. Thus, any review of documents should be specifically delineated as to the nature and extent of our review within our area of specialization."""

suggestions["pattern20"] = """Avoid absolutes (i.e., "all"). Absolutes should be avoided in the context of providing advice. Similarly, different readers may have different interpretations of what these words mean."""

suggestions["pattern21"] = """Implies that we are operating on behalf of our client, taking on client's management responsibilities; possible independence violation."""

suggestions["pattern22"] = """Avoid the phrase when speaking about people as it raises risks relating to various diversity aspects. Equally important is to avoid expressions that relate to age when speaking about a more experienced professional."""

suggestions["pattern23"] = """Avoid taking a position on behalf of the firm - unless you have specifically been engaged to provide an opinion. This is a term defined in professional standards that implies we may be providing a level of assurance. If you are unsure, check with your assigned Risk Management Partner for further clarification."""

suggestions["pattern24"] = """This word/phrase should be avoided in connection with reporting the results of our services as it is ambiguous and often unsubstantiated. 
Avoid absolutes (i.e., "all"). Absolutes should be avoided in the context of providing advice. Similarly, different readers may have different interpretations of what those words mean."""

suggestions["pattern25_1"] = """Implies legal relationship; possible independence vilation and/or a joint business relationship with the client or another party."""
suggestions["pattern25_2"] = """Implies legal relationship; possible independence vilation and/or a joint business relationship with the client or another party."""

suggestions["pattern26"] = """May be interpreted as "off-the-shelf" services; however if we are actually deliverating a firm designed product - please follow the firm's Product & Technology protocols."""

suggestions["pattern27"] = """This word should not be used in isolation when reporting results as it is ambiguous and often unsubstantiated. Our role is not to offer conclusions, rather it is to do the analysis and allow the client to use the information presented to draw their own conclusions. PwC's role should be advisory in nature, we should avoid attest-type terms.
Using the word reasonable in the context of qualitative measures would usually be acceptable provided that the basis for 'reasonableness' is documented within the reporting. However, the word reasonable should not be used in the context of quantitative measures (e.g. account balances, proposed adjustments)."""

suggestions["pattern28_1"] = """We make no implied/actual guarantees.
Should may be used in the event of replacing "must" so not to give a directive or implying an affirmation."""
suggestions["pattern28_2"] = """We make no implied/actual guarantees."""
suggestions["pattern28_3"] = """We make no implied/actual guarantees."""

suggestions["pattern29_1"] = """Most PwC services do not require use of a particular technology tool nor do they convey IP rights or other aspects of software-as-a-service. If you wish to license technology to a client, a separate license agreement will be required.
Please make sure that any technology or product that is going to be included in any materials has gone through the Digital Commercialization process. You can begin this process by submitting a P&T Digital Asset Intake form."""
suggestions["pattern29_2"] = """Most PwC services do not require use of a particular technology tool nor do they convey IP rights or other aspects of software-as-a-service. If you wish to license technology to a client, a separate license agreement will be required.
Please make sure that any technology or product that is going to be included in any materials has gone through the Digital Commercialization process. You can begin this process by submitting a P&T Digital Asset Intake form."""

suggestions["pattern30"] = """Care should be taken when describing our services; may infer a JBR; taking on client's management responsibilities; or implying services that might create an independence concern.
Use of alternative words such as "provide assistance" should also include clarity or specificity as to what we will be assisting with."""

suggestions["pattern31"] = """The word "Tax" is ambiguous so we must be careful to define the nature or type of taxes that will be the subject of our services."""

suggestions["pattern32"] = """We advise/assist; the client implements (e.g., any tax planning/restructuring).
Use of alternative words such as "assist" should also include clarity or specificity as to what we will be assisting with."""

suggestions["pattern33_1"] = """May imply attest services. Scope of services may include the term in the context of what PwC will 'not' perform e.g., "...PwC will not independently validate or verify the information provided"."""
suggestions["pattern33_2"] = """May imply attest services. Scope of services may include the term in the context of what PwC will 'not' perform e.g., "...PwC will not independently validate or verify the information provided"."""

#adding the patterns to the matcher. Future functionality to include adding only certain modules for OFRO to work with
matcher.add("pattern1_1", [pattern1_1])
matcher.add("pattern1_2", [pattern1_2])
matcher.add("pattern1_3", [pattern1_3])
matcher.add("pattern1_4", [pattern1_4])
matcher.add("pattern1_5", [pattern1_5])
matcher.add("pattern1_6", [pattern1_6])

matcher.add("pattern2_1", [pattern2_1])
matcher.add("pattern2_2", [pattern2_2])
matcher.add("pattern2_3", [pattern2_3])
matcher.add("pattern2_4", [pattern2_4])
matcher.add("pattern2_5", [pattern2_5])
matcher.add("pattern2_6", [pattern2_6])
matcher.add("pattern2_7", [pattern2_7])
matcher.add("pattern2_8_1", [pattern2_8_1])
matcher.add("pattern2_8_2", [pattern2_8_2])
matcher.add("pattern2_9_1", [pattern2_9_1])
matcher.add("pattern2_9_2", [pattern2_9_2])
matcher.add("pattern2_10", [pattern2_10])
matcher.add("pattern2_11", [pattern2_11])
matcher.add("pattern2_12", [pattern2_12])

matcher.add("pattern3_1", [pattern3_1])
matcher.add("pattern3_2", [pattern3_2])
matcher.add("pattern3_3", [pattern3_3])
matcher.add("pattern3_4", [pattern3_4])
matcher.add("pattern3_5", [pattern3_5])
matcher.add("pattern3_6", [pattern3_6])
matcher.add("pattern3_7", [pattern3_7])
matcher.add("pattern3_8", [pattern3_8])

matcher.add("pattern4_1", [pattern4_1])
matcher.add("pattern4_2", [pattern4_2])
matcher.add("pattern4_3", [pattern4_3])
matcher.add("pattern4_4", [pattern4_4])
matcher.add("pattern4_5", [pattern4_5])
matcher.add("pattern4_6", [pattern4_6])
matcher.add("pattern4_7", [pattern4_7])
matcher.add("pattern4_8", [pattern4_8])

matcher.add("pattern5_1", [pattern5_1])
matcher.add("pattern5_2", [pattern5_2])
matcher.add("pattern5_3", [pattern5_3])
matcher.add("pattern5_4", [pattern5_4])
matcher.add("pattern5_5", [pattern5_5])
matcher.add("pattern5_6", [pattern5_6])
matcher.add("pattern5_7", [pattern5_7])

matcher.add("pattern6_1", [pattern6_1])

matcher.add("pattern8_1_0", [pattern8_1_0])
matcher.add("pattern8_1_1", [pattern8_1_1])
matcher.add("pattern8_1_2", [pattern8_1_2])
matcher.add("pattern8_2", [pattern8_2])
matcher.add("pattern8_3", [pattern8_3])
matcher.add("pattern8_4", [pattern8_4])
matcher.add("pattern8_4_1", [pattern8_4_1])
matcher.add("pattern8_5", [pattern8_5])
matcher.add("pattern8_5_1", [pattern8_5_1])
matcher.add("pattern8_6", [pattern8_6])
matcher.add("pattern8_7", [pattern8_7])
matcher.add("pattern8_8", [pattern8_8])
matcher.add("pattern8_9", [pattern8_9])
matcher.add("pattern8_10", [pattern8_10])
matcher.add("pattern8_11", [pattern8_11])
matcher.add("pattern8_12", [pattern8_12])
matcher.add("pattern8_13", [pattern8_13])
matcher.add("pattern8_14", [pattern8_14])
matcher.add("pattern8_15", [pattern8_15])

matcher.add("pattern9_1", [pattern9_1])
matcher.add("pattern9_2", [pattern9_2])
matcher.add("pattern9_3", [pattern9_3])

matcher.add("pattern10_1", [pattern10_1])
matcher.add("pattern10_2", [pattern10_2])

matcher.add("pattern11", [pattern11])

matcher.add("pattern12", [pattern12])

matcher.add("pattern13_1", [pattern13_1])
matcher.add("pattern13_2", [pattern13_2])

matcher.add("pattern14_1", [pattern14_1])
matcher.add("pattern14_2", [pattern14_2])

matcher.add("pattern15_1", [pattern15_1])
matcher.add("pattern15_2", [pattern15_2])
matcher.add("pattern15_2_1", [pattern15_2_1])
matcher.add("pattern15_3", [pattern15_3])

matcher.add("pattern16", [pattern16])

matcher.add("pattern17_1", [pattern17_1])
matcher.add("pattern17_2", [pattern17_2])

matcher.add("pattern18_1", [pattern18_1])
matcher.add("pattern18_2", [pattern18_2])
matcher.add("pattern18_3", [pattern18_3])
matcher.add("pattern18_4", [pattern18_4])

matcher.add("pattern19", [pattern19])

matcher.add("pattern20", [pattern20])

matcher.add("pattern21", [pattern21])

matcher.add("pattern22", [pattern22])

matcher.add("pattern23", [pattern23])

matcher.add("pattern24", [pattern24])

matcher.add("pattern25_1", [pattern25_1])
matcher.add("pattern25_2", [pattern25_2])

matcher.add("pattern26", [pattern26])

matcher.add("pattern27", [pattern27])

matcher.add("pattern28_1", [pattern28_1])
matcher.add("pattern28_2", [pattern28_2])
matcher.add("pattern28_3", [pattern28_3])

matcher.add("pattern29_1", [pattern29_1])
matcher.add("pattern29_2", [pattern29_2])

matcher.add("pattern30", [pattern30])

matcher.add("pattern31", [pattern31])

matcher.add("pattern32", [pattern32])

matcher.add("pattern33_1", [pattern33_1])
matcher.add("pattern33_2", [pattern33_2])
#endregion

# Call the matcher on each nlp Doc that each represent a paragraph in the python docx Doc.
docmatches = []
docx_matches = []

#get starting and ending points of each paragraph to add as output for each docx_match. Will be used for isolate_run down the line to understand which paragraph to add the run into

# Calls the matcher on each nlp doc and adds each doc's match outputs into the matches list
for paragraph, text in docs.items():
    docmatches.append(matcher(text, as_spans=True))

# print(docmatches)

# Iterate through the matches found in each nlp Doc, and return them in a format that python docx can understand (by character indices rather than token indices)
for i, matches in enumerate(docmatches):
    for j, span in enumerate(matches):
        docx_matches.append([span.label_, span.text, span.start_char, span.end_char, i])

# for match in docx_matches:
#     print(match)

for index, match in enumerate(docx_matches):
    run = isolate_run(docx.paragraphs[docx_matches[index][4]], docx_matches[index][2], docx_matches[index][3])
    run.add_comment(suggestions[docx_matches[index][0]])

docx.save('Output2.docx')

# for paragraph in docx.paragraphs:
#     for run in paragraph.runs:
#         print(run.text)